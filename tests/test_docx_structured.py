from __future__ import annotations

from docx import Document

from pagespeak.backends._docx_structured import render_markdown

_NUM = """
<w:abstractNum w:abstractNumId="0">
  <w:lvl w:ilvl="0"><w:numFmt w:val="decimal"/></w:lvl>
  <w:lvl w:ilvl="1"><w:numFmt w:val="decimal"/></w:lvl>
</w:abstractNum>
<w:abstractNum w:abstractNumId="9"><w:lvl w:ilvl="0">
  <w:numFmt w:val="bullet"/></w:lvl></w:abstractNum>
<w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>
<w:num w:numId="3"><w:abstractNumId w:val="9"/></w:num>
"""

_NUM2 = """
<w:abstractNum w:abstractNumId="0">
  <w:lvl w:ilvl="0"><w:numFmt w:val="decimal"/></w:lvl>
  <w:lvl w:ilvl="1"><w:numFmt w:val="decimal"/></w:lvl>
</w:abstractNum>
<w:abstractNum w:abstractNumId="1">
  <w:lvl w:ilvl="0"><w:numFmt w:val="decimal"/></w:lvl>
  <w:lvl w:ilvl="1"><w:numFmt w:val="decimal"/></w:lvl>
</w:abstractNum>
<w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>
<w:num w:numId="2"><w:abstractNumId w:val="1"/></w:num>
"""


def _p_list(num_id: int, ilvl: int, text: str) -> str:
    return (
        f'<w:p><w:pPr><w:numPr><w:ilvl w:val="{ilvl}"/>'
        f'<w:numId w:val="{num_id}"/></w:numPr></w:pPr>'
        f"<w:r><w:t>{text}</w:t></w:r></w:p>"
    )


def _p_heading(style: str, text: str) -> str:
    return f'<w:p><w:pPr><w:pStyle w:val="{style}"/></w:pPr><w:r><w:t>{text}</w:t></w:r></w:p>'


def test_heading_style_to_atx(make_docx) -> None:
    # Genuine `Heading N` styles → ATX (no numPr ⇒ literal level).
    # `Title` is NOT a section-heading style (covered by the ilvl0
    # rule when it matters) — see
    # test_title_style_no_numpr_is_not_a_heading.
    xml = _p_heading("Heading1", "Doc") + _p_heading("Heading2", "Sub")
    md = render_markdown(Document(str(make_docx(document_xml=xml))), None)
    assert "# Doc" in md.splitlines()
    assert "## Sub" in md.splitlines()


def test_title_style_no_numpr_is_not_a_heading(make_docx) -> None:
    # A `Title`-styled paragraph with broken numbering (no numPr) ⇒
    # must NOT spuriously promote.
    xml = _p_heading("Heading1", "Real") + _p_heading("Title", "(CHAPTER 2 — Review)")
    lines = render_markdown(Document(str(make_docx(document_xml=xml))), None).splitlines()
    assert "# Real" in lines
    assert "# (CHAPTER 2 — Review)" not in lines
    assert "(CHAPTER 2 — Review)" in lines


def test_non_numpr_heading_style_literal_level(make_docx) -> None:
    xml = _p_heading("Heading2", "DeepOnly")
    md = render_markdown(Document(str(make_docx(document_xml=xml))), None)
    assert "## DeepOnly" in md.splitlines()


def test_numpr_ilvl0_is_heading_regardless_of_style(make_docx) -> None:
    xml = (
        '<w:p><w:pPr><w:pStyle w:val="Heading1"/>'
        '<w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr>'
        "</w:pPr><w:r><w:t>Chap</w:t></w:r></w:p>"
        "<w:p><w:r><w:t>body content under the chapter</w:t></w:r></w:p>"
    )
    md = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM))),
        None,
        outline_heading_depth=1,
    )
    assert "# Chap" in md.splitlines()


def test_pstyle_on_outline_node_is_ignored_ilvl_is_signal(make_docx) -> None:
    # `pStyle` (Title/Heading1) on a numPr outline paragraph is NOISE.
    # The signal is `ilvl` + outline_heading_depth (default 1):
    # ilvl0 → `#`; ilvl≥1 → the RETAINED nested outline list. A
    # Heading1-styled ilvl1 node is a list item, NOT a `##`.
    body = "<w:p><w:r><w:t>body prose so the section is not bodyless</w:t></w:r></w:p>"
    h1 = (
        '<w:p><w:pPr><w:pStyle w:val="Heading1"/>'
        '<w:numPr><w:ilvl w:val="1"/><w:numId w:val="1"/></w:numPr>'
        "</w:pPr><w:r><w:t>child</w:t></w:r></w:p>"
    )
    xml = _p_list(1, 0, "S1") + body + h1 + body + _p_list(1, 0, "S2") + body
    lines = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM))),
        None,
        outline_heading_depth=1,
    ).splitlines()
    assert "# S1" in lines  # ilvl0 section
    # Heading1@ilvl1 → a retained outline list item (pStyle ignored);
    # the trailing non-numbered body appends inline to its line.
    assert any(ln.startswith("1. child") for ln in lines)
    assert not any(ln.startswith("## child") for ln in lines)  # pStyle ignored
    assert not any(ln.startswith("# child") for ln in lines)


def test_ordered_list_running_numbers_and_reset(make_docx) -> None:
    # depth=1 mode: ilvl0 → `#` section; ilvl≥1 → RETAINED nested
    # outline list. Outline numbering restarts under each ilvl0
    # section; the first list level sits at column 0.
    xml = (
        _p_list(1, 0, "A")
        + _p_list(1, 1, "a1")
        + _p_list(1, 1, "a2")
        + _p_list(1, 0, "B")
        + _p_list(1, 1, "b1")
    )
    md = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM))),
        None,
        outline_heading_depth=1,
    ).splitlines()
    assert "# A" in md
    assert "1. a1" in md
    assert "2. a2" in md
    assert "# B" in md
    assert "1. b1" in md  # counter reset under the new ilvl0 section
    assert "  1. a1" not in md  # first list level at column 0


def test_outline_heading_depth_switch(make_docx) -> None:
    # The switch: how many top outline levels promote to headings.
    xml = _p_list(1, 0, "S") + _p_list(1, 1, "a") + _p_list(1, 2, "b")
    doc = lambda: Document(str(make_docx(document_xml=xml, numbering_xml=_NUM)))  # noqa: E731

    # depth=0 → whole outline retained as a list (ilvl0 also a list).
    d0 = render_markdown(doc(), None, outline_heading_depth=0).splitlines()
    assert "1. S" in d0
    assert "# S" not in d0

    # depth=1 (default) → ilvl0 `#`, ilvl≥1 retained list.
    d1 = render_markdown(doc(), None, outline_heading_depth=1).splitlines()
    assert "# S" in d1
    assert "1. a" in d1
    assert "    1. b" in d1

    # depth=2 → ilvl0 `#`, ilvl1 `##`, ilvl≥2 retained list at col 0.
    d2 = render_markdown(doc(), None, outline_heading_depth=2).splitlines()
    assert "# S" in d2
    assert "## a" in d2
    assert "1. b" in d2  # ilvl2 → first retained-list level, column 0
    assert "    1. b" not in d2


def test_independent_numids_dont_share_counter(make_docx) -> None:
    # Each numId keeps its OWN counter. These synthetic paragraphs
    # carry NO `w:ind` and `_NUM2` defines no level indent, so depth
    # falls back to the ilvl ½-inch ladder — numId1-ilvl1 and
    # numId2-ilvl1 resolve to the same indent ⇒ same depth (siblings).
    # (Real docs carry explicit indents — see
    # test_indent_drives_nesting_depth_across_numids.)
    xml = (
        _p_list(1, 0, "Sec")
        + _p_list(1, 1, "n1-a")
        + _p_list(2, 1, "n2-a")
        + _p_list(1, 1, "n1-b")
        + _p_list(2, 1, "n2-b")
    )
    md = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM2))),
        None,
        outline_heading_depth=1,
    ).splitlines()
    assert "# Sec" in md  # ilvl0 → #
    assert "1. n1-a" in md
    assert "1. n2-a" in md  # different numId → own counter
    assert "2. n1-b" in md
    assert "2. n2-b" in md


def test_empty_heading_styled_spacer_does_not_denest_outline(make_docx) -> None:
    # An EMPTY `Heading 1`-styled paragraph (a Word spacer) sitting
    # between outline items must be a no-op — it must NOT reset the
    # outline, or the next item collapses to column 0 (a sibling item
    # knocked to the main level).
    spacer = '<w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr><w:r><w:t></w:t></w:r></w:p>'
    xml = (
        _p_list(1, 0, "Sec")
        + _p_list(1, 1, "child")
        + _p_list(1, 2, "grandchild")
        + spacer
        + _p_list(1, 1, "sibling-of-child")
    )
    lines = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM))), None
    ).splitlines()
    assert "1. Sec" in lines  # ilvl0 → depth 0
    assert "    1. child" in lines  # ilvl1 → depth 1
    assert "        1. grandchild" in lines  # ilvl2 → depth 2
    # after the empty Heading-styled spacer the outline is intact:
    # the next ilvl1 item keeps depth 1, NOT collapsed to column 0.
    assert "    2. sibling-of-child" in lines
    assert "2. sibling-of-child" not in lines
    assert "1. sibling-of-child" not in lines
    # the empty spacer itself emitted nothing (no bare "#").
    assert "#" not in lines
    assert "# " not in lines


def test_indent_drives_nesting_depth_across_numids(make_docx) -> None:
    # Word lays the outline out by RESOLVED LEFT INDENT, not
    # (numId, ilvl). A foreign bullet numId whose own ilvl is 0 but
    # whose paragraph `w:ind` places it deep must nest at that visual
    # depth — never collapse to column 0.
    def p_ind(num_id: int, ilvl: int, left: int, text: str) -> str:
        return (
            f'<w:p><w:pPr><w:ind w:left="{left}"/>'
            f'<w:numPr><w:ilvl w:val="{ilvl}"/>'
            f'<w:numId w:val="{num_id}"/></w:numPr></w:pPr>'
            f"<w:r><w:t>{text}</w:t></w:r></w:p>"
        )

    xml = (
        p_ind(1, 0, 720, "A")  # left 720  → depth 0
        + p_ind(1, 1, 1440, "a1")  # left 1440 → depth 1
        + p_ind(1, 2, 2160, "a1x")  # left 2160 → depth 2
        + p_ind(5, 0, 2160, "bullet-same")  # numId5 ilvl0, left 2160 → depth 2 (sibling, NOT col 0)
        + p_ind(2, 0, 2880, "bullet-deeper")  # numId2 ilvl0, left 2880 → depth 3
        + p_ind(1, 3, 1440, "back-up")  # left 1440 → depth 1 (back to numId1)
    )
    # default depth=0: the WHOLE outline is a retained list; depth is
    # driven by left indent.
    lines = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM2))), None
    ).splitlines()
    assert "1. A" in lines  # left 720  → depth 0 (col 0)
    assert "    1. a1" in lines  # left 1440 → depth 1
    assert "        1. a1x" in lines  # left 2160 → depth 2
    # foreign numId5 ilvl0 at left 2160 → SAME depth as a1x (NOT col 0)
    assert "        1. bullet-same" in lines
    assert "            1. bullet-deeper" in lines  # numId2 ilvl0, left 2880 → depth 3
    assert "    1. back-up" in lines  # left 1440 → back to the a1 depth


def test_runs_bold_italic(make_docx) -> None:
    xml = (
        "<w:p><w:r><w:rPr><w:b/></w:rPr><w:t>bold</w:t></w:r>"
        "<w:r><w:t> plain </w:t></w:r>"
        "<w:r><w:rPr><w:i/></w:rPr><w:t>it</w:t></w:r></w:p>"
    )
    md = render_markdown(Document(str(make_docx(document_xml=xml))), None)
    assert "**bold** plain *it*" in md


def test_inline_image_emitted_positionally(make_docx, tmp_path) -> None:
    png = (
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00"
        b"\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx"
        b"\x9cc\x00\x01\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND"
        b"\xaeB`\x82"
    )
    drawing = (
        "<w:p><w:r><w:drawing><wp:inline>"
        '<wp:docPr id="1" name="pic" descr="A widget"/>'
        "<a:graphic><a:graphicData><pic:pic "
        'xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:blipFill><a:blip r:embed="rIdImg"/></pic:blipFill>'
        "</pic:pic></a:graphicData></a:graphic>"
        "</wp:inline></w:drawing></w:r></w:p>"
    )
    path = make_docx(
        document_xml="<w:p><w:r><w:t>before</w:t></w:r></w:p>" + drawing,
        extra_parts={"word/media/image1.png": png},
        doc_rels=(
            '<Relationship Id="rIdImg" Type="http://schemas.'
            "openxmlformats.org/officeDocument/2006/relationships/image"
            '" Target="media/image1.png"/>'
        ),
    )
    md = render_markdown(Document(str(path)), tmp_path)
    assert "before" in md
    assert "![A widget](images/image1.png)" in md
    assert (tmp_path / "images" / "image1.png").read_bytes() == png


def test_convert_structured_returns_ingestresult(make_docx, tmp_path) -> None:
    from pagespeak.backends._docx_structured import convert_structured

    path = make_docx(document_xml="<w:p><w:r><w:t>Hi</w:t></w:r></w:p>")
    res = convert_structured(path, output_dir=tmp_path)
    assert res.source_format == "docx"
    assert res.markdown.strip() == "Hi"


def test_convert_structured_empty_doc_empty_markdown(make_docx, tmp_path) -> None:
    from pagespeak.backends._docx_structured import convert_structured

    res = convert_structured(make_docx(document_xml=""), output_dir=tmp_path)
    assert res.markdown == ""


def test_convert_structured_falls_back_on_parse_error(tmp_path) -> None:
    from unittest.mock import patch

    from pagespeak.backends._docx_structured import convert_structured
    from pagespeak.models._models import IngestResult

    bad = tmp_path / "bad.docx"
    bad.write_bytes(b"not a zip")
    sentinel = IngestResult(markdown="FALLBACK", source_format="docx")
    with patch(
        "pagespeak.backends._docx.convert_with_markitdown",
        return_value=sentinel,
    ) as mk:
        res = convert_structured(bad, output_dir=tmp_path)
    mk.assert_called_once()
    assert res.markdown == "FALLBACK"


def test_pstyle_ignored_ilvl_drives_retained_outline_2topics_shape(make_docx) -> None:
    def styled(style, ilvl, text):  # noqa: ANN001,ANN202
        return (
            f'<w:p><w:pPr><w:pStyle w:val="{style}"/>'
            f'<w:numPr><w:ilvl w:val="{ilvl}"/>'
            f'<w:numId w:val="1"/></w:numPr></w:pPr>'
            f"<w:r><w:t>{text}</w:t></w:r></w:p>"
        )

    xml = (
        styled("Title", 0, "Definition of a widget")
        + styled("Title", 1, "Parts")
        + styled("normal", 2, "Frame")
        + styled("normal", 1, "Definition of assembly")
        + styled("Title", 0, "Four basic types")
        + styled("Heading1", 1, "surface layer")
        + styled("Title", 0, "Surfaces")
        + styled("Heading1", 1, "Surface is found")
        + styled("Title", 2, "covering outer faces")
    )
    lines = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM))),
        None,
        outline_heading_depth=1,
    ).splitlines()
    # `pStyle` (Title/Heading1/normal) is NOISE and ignored; `ilvl` +
    # depth=1 is the signal. ilvl0 → `#` section; ilvl≥1 → the
    # RETAINED nested outline list (numbering restarts per ilvl0
    # section; deeper ilvl indents one 4-space level).
    assert "# Definition of a widget" in lines  # ilvl0
    assert "1. Parts" in lines  # ilvl1 → outline list, col 0
    assert "    1. Frame" in lines  # ilvl2 → nested under Parts
    assert "2. Definition of assembly" in lines  # ilvl1 sibling
    assert "# Four basic types" in lines  # ilvl0 (counter resets)
    assert "1. surface layer" in lines  # Heading1 ignored → list
    assert "# Surfaces" in lines  # ilvl0
    assert "1. Surface is found" in lines  # Heading1 ignored → list
    assert "    1. covering outer faces" in lines  # ilvl2 nested
    # never a heading from ilvl≥1, never wrong
    assert "## Parts" not in lines
    assert "# Parts" not in lines
    assert "## surface layer" not in lines


def test_inline_continuation_folds_into_list_item_keeps_outline(make_docx) -> None:
    # A non-numbered paragraph (caption / note / reaction label /
    # resource link / inline figure) BETWEEN numbered outline items is
    # appended INLINE to the current list item's line with a single
    # space — NO hard return (no blank line, no `<br>`) before or
    # after. No separate block ⇒ no renderer can turn it into a code
    # block or de-nest the outline that follows. Every outline item
    # stays exactly one source line.
    xml = (
        _p_list(1, 0, "Widget")
        + _p_list(1, 1, "Ports")
        + _p_list(1, 2, "spare port detail")
        + "<w:p><w:r><w:t>Figure 1. Widget interface components</w:t></w:r></w:p>"
        + _p_list(1, 1, "Transfer")
        + _p_list(1, 2, "passive routing")
    )
    lines = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM))), None
    ).splitlines()
    # caption appended inline (single space, no hard return); never a
    # standalone line, no `<br>`.
    assert "        1. spare port detail Figure 1. Widget interface components" in lines
    assert "Figure 1. Widget interface components" not in lines
    assert not any("<br>" in ln for ln in lines)
    # "Transfer" keeps its outline level (ilvl1 → 4-space): not
    # de-nested to column 0, not a code block.
    assert "    2. Transfer" in lines
    assert "2. Transfer" not in lines
    assert "        1. passive routing" in lines  # child nests one deeper


def test_blank_line_before_list_after_plain_paragraph(make_docx) -> None:
    # A retained-outline list immediately after a body paragraph with
    # NO blank line is folded into it as run-on prose by real
    # renderers. A separating blank line must precede the list.
    xml = (
        '<w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr>'
        "<w:r><w:t>Topics</w:t></w:r></w:p>"
        "<w:p><w:r><w:t>Intro sentence.</w:t></w:r></w:p>"
        '<w:p><w:pPr><w:numPr><w:ilvl w:val="1"/>'
        '<w:numId w:val="1"/></w:numPr></w:pPr>'
        "<w:r><w:t>First item</w:t></w:r></w:p>"
    )
    lines = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM))), None
    ).splitlines()
    i = lines.index("Intro sentence.")
    assert lines[i + 1] == ""  # blank separates paragraph from list
    assert lines[i + 2] == "1. First item"  # numPr ilvl1 → outline list, col 0


def test_numpr_ilvl0_boilerplate_preamble_is_a_heading(make_docx) -> None:
    # `I. Before you begin, ...:` is numPr ilvl0 — the document's
    # top-level spine, parent of the A/B/C sub-outline. A text-pattern
    # boilerplate-junk rule must NOT pull a numPr structural node out
    # of the outline.
    xml = (
        '<w:p><w:pPr><w:pStyle w:val="Title"/>'
        '<w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr></w:pPr>'
        "<w:r><w:t>Before you begin, make sure you have "
        "mastered the following topics: </w:t></w:r></w:p>"
        '<w:p><w:pPr><w:numPr><w:ilvl w:val="1"/>'
        '<w:numId w:val="1"/></w:numPr></w:pPr>'
        "<w:r><w:t>Basic terms</w:t></w:r></w:p>"
    )
    lines = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM))),
        None,
        outline_heading_depth=1,
    ).splitlines()
    assert "# Before you begin, make sure you have mastered the following topics:" in lines
    # not demoted to a stray body paragraph
    assert "Before you begin, make sure you have mastered the following topics:" not in lines


def test_blank_line_before_heading_after_list_item(make_docx) -> None:
    # A heading (ilvl0) directly after a retained-outline list item
    # (ilvl1) with no blank line is folded into it by some markdown
    # renderers ("…wireless # Some types…"). A blank must separate them.
    body = "<w:p><w:r><w:t>section body</w:t></w:r></w:p>"
    xml = (
        '<w:p><w:pPr><w:numPr><w:ilvl w:val="1"/>'
        '<w:numId w:val="1"/></w:numPr></w:pPr>'
        "<w:r><w:t>Wired v. wireless</w:t></w:r></w:p>"
        '<w:p><w:pPr><w:numPr><w:ilvl w:val="0"/>'
        '<w:numId w:val="1"/></w:numPr></w:pPr>'
        "<w:r><w:t>Some types of signal routing</w:t></w:r></w:p>" + body
    )
    lines = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM))),
        None,
        outline_heading_depth=1,
    ).splitlines()
    assert "1. Wired v. wireless" in lines  # ilvl1 → retained list item
    h = lines.index("# Some types of signal routing")  # ilvl0 → heading
    assert lines[h - 1] == ""  # blank separates list item from heading


def test_adjacent_same_format_runs_are_coalesced(make_docx) -> None:
    # Word stores one visual token (ID7) across several adjacent bold
    # runs, each with its own w:rPr. Wrapping each independently shatters
    # it into `**ID****7**** + K**`. Coalesce same-format neighbours.
    xml = (
        "<w:p>"
        "<w:r><w:rPr><w:b/></w:rPr><w:t>ID</w:t></w:r>"
        "<w:r><w:rPr><w:b/></w:rPr><w:t>7</w:t></w:r>"
        '<w:r><w:rPr><w:b/></w:rPr><w:t xml:space="preserve"> + K</w:t></w:r>'
        "</w:p>"
    )
    md = render_markdown(Document(str(make_docx(document_xml=xml))), None)
    assert "**ID7 + K**" in md
    assert "****" not in md


def test_empty_text_run_between_same_format_runs_drops_no_markers(make_docx) -> None:
    xml = (
        "<w:p>"
        "<w:r><w:rPr><w:b/></w:rPr><w:t>Base</w:t></w:r>"
        "<w:r><w:rPr><w:b/></w:rPr><w:t></w:t></w:r>"
        "<w:r><w:rPr><w:b/></w:rPr><w:t>line</w:t></w:r>"
        "</w:p>"
    )
    md = render_markdown(Document(str(make_docx(document_xml=xml))), None)
    assert "**Baseline**" in md
    assert "****" not in md


def test_hyperlink_remains_a_segment_boundary(make_docx) -> None:
    xml = (
        "<w:p>"
        "<w:r><w:rPr><w:b/></w:rPr><w:t>see </w:t></w:r>"
        '<w:hyperlink r:id="rIdL"><w:r><w:t>here</w:t></w:r></w:hyperlink>'
        "<w:r><w:rPr><w:b/></w:rPr><w:t> now</w:t></w:r>"
        "</w:p>"
    )
    path = make_docx(
        document_xml=xml,
        doc_rels=(
            '<Relationship Id="rIdL" Type="http://schemas.openxmlformats.'
            'org/officeDocument/2006/relationships/hyperlink" '
            'Target="https://x.test/" TargetMode="External"/>'
        ),
    )
    md = render_markdown(Document(str(path)), None)
    assert "**see **[here](https://x.test/)** now**" in md


def test_table_renders_gfm_at_document_position(make_docx) -> None:
    grid = "<w:tblGrid><w:gridCol/><w:gridCol/></w:tblGrid>"
    tbl = (
        f"<w:tbl>{grid}"
        "<w:tr><w:tc><w:p><w:r><w:t>H1</w:t></w:r></w:p></w:tc>"
        "<w:tc><w:p><w:r><w:t>H2</w:t></w:r></w:p></w:tc></w:tr>"
        "<w:tr><w:tc><w:p><w:r><w:t>a</w:t></w:r></w:p></w:tc>"
        "<w:tc><w:p><w:r><w:t>b</w:t></w:r></w:p></w:tc></w:tr>"
        "</w:tbl>"
    )
    xml = "<w:p><w:r><w:t>before</w:t></w:r></w:p>" + tbl
    md = render_markdown(Document(str(make_docx(document_xml=xml))), None)
    lines = md.splitlines()
    assert "before" in lines
    assert "| H1 | H2 |" in lines
    assert "| --- | --- |" in lines
    assert "| a | b |" in lines
    assert "<!-- TABLE" not in md
    assert lines.index("| H1 | H2 |") > lines.index("before")


def test_table_between_outline_items_keeps_them_parallel(make_docx) -> None:
    # A table is CONTENT inside the outline, not a section break. An
    # outline item after a table must stay at its true depth — parallel
    # to its same-level sibling before the table — NOT reset to column 0
    # (`Optical controls` ∥ `Manual controls`).
    grid = "<w:tblGrid><w:gridCol/></w:tblGrid>"
    tbl = f"<w:tbl>{grid}<w:tr><w:tc><w:p><w:r><w:t>c</w:t></w:r></w:p></w:tc></w:tr></w:tbl>"
    xml = (
        _p_list(1, 0, "Section")
        + _p_list(1, 1, "Manual controls")
        + _p_list(1, 2, "detail")
        + tbl
        + _p_list(1, 1, "Optical controls")
    )
    lines = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM))), None
    ).splitlines()
    assert "1. Section" in lines  # ilvl0 → depth 0
    assert "    1. Manual controls" in lines  # ilvl1 → depth 1
    assert "        1. detail" in lines  # ilvl2 → depth 2
    assert "| c |" in lines  # the table renders
    # after the table the outline is intact: the next ilvl1 item is
    # parallel to "Manual controls" (depth 1), NOT de-nested.
    assert "    2. Optical controls" in lines
    assert "1. Optical controls" not in lines
    assert "2. Optical controls" not in lines


def test_ilvl0_and_styled_heading_emphasis_stripped(make_docx) -> None:
    xml = (
        '<w:p><w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/>'
        "</w:numPr></w:pPr>"
        "<w:r><w:rPr><w:b/></w:rPr><w:t>Cooling</w:t></w:r></w:p>"
        '<w:p><w:pPr><w:pStyle w:val="Heading2"/></w:pPr>'
        "<w:r><w:rPr><w:b/></w:rPr><w:t>Heat </w:t></w:r>"
        "<w:r><w:rPr><w:b/></w:rPr><w:t>exchange</w:t></w:r></w:p>"
    )
    md = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM))),
        None,
        outline_heading_depth=1,
    )
    lines = md.splitlines()
    assert "# Cooling" in lines
    assert "## Heat exchange" in lines
    assert "**" not in md


def test_headless_bold_body_title_promoted_when_doc_has_heading(make_docx) -> None:
    xml = (
        "<w:p><w:r><w:rPr><w:b/></w:rPr>"
        "<w:t>Acoustic Widgetry (Chapters 16, 17)</w:t></w:r></w:p>"
        + _p_list(1, 0, "Steps of operation")
        + _p_list(1, 1, "Toolcraft of the signal path")
    )
    lines = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM))),
        None,
        outline_heading_depth=1,
    ).splitlines()
    assert "# Acoustic Widgetry (Chapters 16, 17)" in lines
    assert "# Steps of operation" in lines
    assert "Acoustic Widgetry (Chapters 16, 17)" not in lines


def test_lone_body_paragraph_not_promoted_without_heading(make_docx) -> None:
    xml = "<w:p><w:r><w:rPr><w:b/></w:rPr><w:t>Just a bold line</w:t></w:r></w:p>"
    md = render_markdown(Document(str(make_docx(document_xml=xml))), None)
    assert md.strip() == "**Just a bold line**"


def test_numpr_first_heading_blocks_body_title_promotion(make_docx) -> None:
    xml = _p_list(1, 0, "Intro") + "<w:p><w:r><w:t>Plain body sentence.</w:t></w:r></w:p>"
    lines = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM))),
        None,
        outline_heading_depth=1,
    ).splitlines()
    assert "# Intro" in lines
    assert "Plain body sentence." in lines
    assert "# Plain body sentence." not in lines


def test_ilvl0_sentence_shaped_node_is_a_faithful_heading(make_docx) -> None:
    # De-heuristic contract: a numPr ilvl0 node is a heading because
    # Word's STRUCTURE says so — never demoted by its wording. A
    # sentence-shaped ilvl0 title WITH real body under it stays `#`.
    # (A bodyless run is still handled by the STRUCTURAL
    # demote_nonsection_h1 — covered in test_docx_quality.py.)
    body = "<w:p><w:r><w:t>real body prose under the section</w:t></w:r></w:p>"
    xml = (
        _p_list(1, 0, "Cooling and airflow mechanics")
        + body
        + _p_list(1, 0, "You will find it on the surface of the panel.")
        + body
        + _p_list(1, 0, "In standby mode, when power is present, the following occurs:")
        + body
    )
    lines = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM))),
        None,
        outline_heading_depth=1,
    ).splitlines()
    assert "# Cooling and airflow mechanics" in lines
    # Sentence / `:`-lead-in wording is NOT demoted — faithful to the
    # author's outline structure.
    assert "# You will find it on the surface of the panel." in lines
    assert "# In standby mode, when power is present, the following occurs:" in lines


def test_render_demotes_bodyless_label_run_keeps_real_section(make_docx) -> None:
    xml = (
        "<w:p><w:r><w:rPr><w:b/></w:rPr>"
        "<w:t>Optical Widgetry (Chapters 18, 19)</w:t></w:r></w:p>"
        + _p_list(1, 0, "NDX")
        + _p_list(1, 0, "Capacitor")
        + _p_list(1, 0, "Optical functions")
        + _p_list(1, 1, "Process input stream")
    )
    lines = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM))),
        None,
        outline_heading_depth=1,
    ).splitlines()
    assert "# Optical Widgetry (Chapters 18, 19)" in lines  # title protected
    assert "NDX" in lines
    assert "# NDX" not in lines  # bodyless ilvl0 → demoted (structural)
    assert "Capacitor" in lines
    assert "# Capacitor" not in lines
    assert "# Optical functions" in lines  # has an outline-list body → real
    assert "1. Process input stream" in lines  # ilvl1 → retained list item


def test_ilvl0_section_with_outline_list_body_is_kept(make_docx) -> None:
    # An ilvl0 section whose body is the retained outline list (first
    # item `1.`) is a REAL section — demote_nonsection_h1 only demotes
    # a bodyless `#` or one whose first body item is a `>=2.`
    # continuation. A `1.` first item means the section owns its list.
    xml = (
        "<w:p><w:r><w:rPr><w:b/></w:rPr><w:t>Widget Assembly</w:t></w:r></w:p>"
        + _p_list(1, 0, "Real Sec")
        + _p_list(1, 1, "a1")
        + _p_list(2, 1, "outer1")
        + _p_list(1, 0, "modules submodules")
        + _p_list(2, 1, "outer2")
    )
    lines = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM2))),
        None,
        outline_heading_depth=1,
    ).splitlines()
    assert "# Widget Assembly" in lines  # promoted title, protected
    assert "# Real Sec" in lines  # first body item `1. a1` → real section
    assert "1. a1" in lines
    # synthetic: no `w:ind`, no level indent → ilvl ladder fallback,
    # so numId2-ilvl1 resolves to numId1-ilvl1's depth (sibling).
    assert "1. outer1" in lines
    assert "# modules submodules" in lines  # first body item `1. outer2` → real
    assert "1. outer2" in lines  # counter reset after the ilvl0 heading
    assert "## a1" not in lines


def test_empty_ilvl0_heading_skipped(make_docx) -> None:
    xml = (
        _p_list(1, 0, "Real Section")
        + _p_list(1, 1, "real body content")
        + '<w:p><w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/>'
        "</w:numPr></w:pPr><w:r><w:t></w:t></w:r></w:p>"
    )
    lines = render_markdown(
        Document(str(make_docx(document_xml=xml, numbering_xml=_NUM))),
        None,
        outline_heading_depth=1,
    ).splitlines()
    assert "# Real Section" in lines
    assert "#" not in lines
    assert "# " not in lines
