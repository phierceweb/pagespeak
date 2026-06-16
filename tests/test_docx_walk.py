from __future__ import annotations

from docx import Document

from pagespeak.backends._docx_walk import build_numfmt_map, iter_body


def test_make_docx_builds_openable_file(make_docx) -> None:
    path = make_docx(
        document_xml=("<w:p><w:r><w:t>Hello</w:t></w:r></w:p>"),
        numbering_xml="",
    )
    doc = Document(str(path))
    assert [p.text for p in doc.paragraphs] == ["Hello"]


_NUMBERING = """
<w:abstractNum w:abstractNumId="0">
  <w:lvl w:ilvl="0"><w:numFmt w:val="decimal"/></w:lvl>
  <w:lvl w:ilvl="1"><w:numFmt w:val="lowerLetter"/></w:lvl>
</w:abstractNum>
<w:abstractNum w:abstractNumId="5">
  <w:lvl w:ilvl="0"><w:numFmt w:val="bullet"/></w:lvl>
</w:abstractNum>
<w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>
<w:num w:numId="7"><w:abstractNumId w:val="5"/></w:num>
"""


def test_build_numfmt_map_resolves_ordered_and_bullet(make_docx) -> None:
    path = make_docx(
        document_xml="<w:p><w:r><w:t>x</w:t></w:r></w:p>",
        numbering_xml=_NUMBERING,
    )
    m = build_numfmt_map(Document(str(path)))
    assert m[(1, 0)] == "decimal"
    assert m[(1, 1)] == "lowerLetter"
    assert m[(7, 0)] == "bullet"


def test_build_numfmt_map_missing_numbering_part_is_empty(make_docx) -> None:
    path = make_docx(document_xml="<w:p><w:r><w:t>x</w:t></w:r></w:p>")
    assert build_numfmt_map(Document(str(path))) == {}


def test_iter_body_preserves_paragraph_table_order(make_docx) -> None:
    doc_xml = (
        "<w:p><w:r><w:t>before</w:t></w:r></w:p>"
        "<w:tbl><w:tr><w:tc><w:p><w:r><w:t>cell</w:t></w:r></w:p>"
        "</w:tc></w:tr></w:tbl>"
        "<w:p><w:r><w:t>after</w:t></w:r></w:p>"
    )
    items = list(iter_body(Document(str(make_docx(document_xml=doc_xml)))))
    assert [i.kind for i in items] == ["paragraph", "table", "paragraph"]
    assert items[0].obj.text == "before"
    assert items[2].obj.text == "after"
