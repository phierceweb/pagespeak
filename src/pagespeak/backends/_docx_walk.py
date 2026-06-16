"""Body-order traversal + Word numbering resolution for the
structure-faithful DOCX backend.

Two concerns, both isolated here for unit testing:
- `build_numfmt_map`: resolve every (numId, ilvl) to its w:numFmt
  ("bullet" vs an ordered format like "decimal") by walking
  word/numbering.xml. Missing/odd entries default to "decimal".
- `iter_body`: yield body children (paragraphs, tables) in true
  document order, which python-docx's `.paragraphs` does not preserve.
"""

from __future__ import annotations

from collections.abc import Iterator
from dataclasses import dataclass
from typing import Any

from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

ORDERED_DEFAULT = "decimal"


def build_numfmt_map(document: Any) -> dict[tuple[int, int], str]:
    """Map (numId, ilvl) -> numFmt string. "bullet" means unordered;
    any other value (decimal/lowerLetter/lowerRoman/...) is ordered.
    Empty dict when the document has no numbering part."""
    try:
        numbering_part = document.part.numbering_part
    except (NotImplementedError, KeyError, AttributeError):
        return {}
    if numbering_part is None:
        return {}
    root = numbering_part.element

    abstract: dict[str, dict[int, str]] = {}
    for anum in root.findall(qn("w:abstractNum")):
        aid = anum.get(qn("w:abstractNumId"))
        if aid is None:
            continue
        levels: dict[int, str] = {}
        for lvl in anum.findall(qn("w:lvl")):
            ilvl_raw = lvl.get(qn("w:ilvl"))
            fmt_el = lvl.find(qn("w:numFmt"))
            if ilvl_raw is None or fmt_el is None:
                continue
            val = fmt_el.get(qn("w:val")) or ORDERED_DEFAULT
            levels[int(ilvl_raw)] = val
        abstract[aid] = levels

    out: dict[tuple[int, int], str] = {}
    for num in root.findall(qn("w:num")):
        nid_raw = num.get(qn("w:numId"))
        aref = num.find(qn("w:abstractNumId"))
        if nid_raw is None or aref is None:
            continue
        aid = aref.get(qn("w:val"))
        nid = int(nid_raw)
        for ilvl, fmt in abstract.get(aid, {}).items():
            out[(nid, ilvl)] = fmt
    return out


def build_numindent_map(document: Any) -> dict[tuple[int, int], int]:
    """Map (numId, ilvl) -> the numbering level's left indent (twips),
    read from ``w:lvl/w:pPr/w:ind/@w:left`` in word/numbering.xml.

    Word lays out the *visual* outline by resolved left indent, not by
    ``(numId, ilvl)`` (ilvl is per-numId; a separate bullet numbering
    restarts at ilvl0 yet is visually deep). This is the numbering-
    level fallback used when a paragraph carries no direct ``w:ind``.
    Missing entries are simply absent (the caller falls back further)."""
    try:
        numbering_part = document.part.numbering_part
    except (NotImplementedError, KeyError, AttributeError):
        return {}
    if numbering_part is None:
        return {}
    root = numbering_part.element

    abstract: dict[str, dict[int, int]] = {}
    for anum in root.findall(qn("w:abstractNum")):
        aid = anum.get(qn("w:abstractNumId"))
        if aid is None:
            continue
        levels: dict[int, int] = {}
        for lvl in anum.findall(qn("w:lvl")):
            ilvl_raw = lvl.get(qn("w:ilvl"))
            ppr = lvl.find(qn("w:pPr"))
            ind = ppr.find(qn("w:ind")) if ppr is not None else None
            left = ind.get(qn("w:left")) if ind is not None else None
            if ilvl_raw is None or left is None:
                continue
            try:
                levels[int(ilvl_raw)] = int(left)
            except ValueError:
                continue
        abstract[aid] = levels

    out: dict[tuple[int, int], int] = {}
    for num in root.findall(qn("w:num")):
        nid_raw = num.get(qn("w:numId"))
        aref = num.find(qn("w:abstractNumId"))
        if nid_raw is None or aref is None:
            continue
        aid = aref.get(qn("w:val"))
        nid = int(nid_raw)
        for ilvl, left in abstract.get(aid, {}).items():
            out[(nid, ilvl)] = left
    return out


@dataclass(frozen=True)
class BodyItem:
    """One top-level body child in document order."""

    kind: str  # "paragraph" | "table"
    obj: Any  # docx Paragraph | Table


def iter_body(document: Any) -> Iterator[BodyItem]:
    """Yield paragraphs and tables in true document order. (python-docx
    `.paragraphs` / `.tables` each lose the interleaving.)"""
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield BodyItem("paragraph", Paragraph(child, document))
        elif child.tag == qn("w:tbl"):
            yield BodyItem("table", Table(child, document))
        # sectPr / bookmarks / other -> skipped (not body content)
